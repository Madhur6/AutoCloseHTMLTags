from django.shortcuts import render, redirect
from django.contrib import messages
from .models import Post

from django.http import HttpResponse

from django import forms

from markdown2 import Markdown

from docx import Document
from io import BytesIO


class PostForm(forms.ModelForm):
    class Meta:
        model = Post
        fields = ['title', 'content']
        widgets = {
            'title': forms.TextInput(attrs={'autofocus':'autofocus', 'autocomplete':'off', 'class':'form-control', 'id':'id_title', 'placeholder':'Doc title...'}),
            'content': forms.Textarea(attrs={'rows':15, 'cols':20, 'autocomplete':'off', 'class':'form-control', 'id':'id_content', 'placeholder':'start writing...'})
        }

# Create your views here.
def index(request):
    if request.method == "POST":
        form = PostForm(request.POST)
        if form.is_valid():
            title = form.cleaned_data['title']
            content = form.cleaned_data['content']

            post = Post.objects.create(title=title, content=content)
            post.save()

            messages.info(request, 'Post created successfully!')
            return redirect('newapp:index')
        else:
            return render(request, "app1/index.html", {"form":form})
        
    return render(request, 'app1/index.html', {
        'form': PostForm()
    })



def CheckPost(request, post_id):
    post = Post.objects.get(pk=post_id)
    return render(request, "app1/post.html", {"post":post,"post_content":markdown(post.content)})


def display(request):
    posts = Post.objects.all()
    return render(request, "app1/display.html", {
        "posts": posts
    })


def edit_post(request, post_id):
    post = Post.objects.get(pk=post_id)

    if request.method == "POST":
        form = PostForm(request.POST, instance=post)
        if form.is_valid():
            form.save()
            messages.success(request, 'Post updated successfully!')
            return redirect('newapp:CheckPost', post_id=post_id)
    else:
        form = PostForm(instance=post, initial={'content': post.content})

    return render(request, 'app1/index.html', {'form': form, 'post': post})



def download_word_doc(request, post_id):
    post = Post.objects.get(pk=post_id)
    markdown_content = post.content

    # Create a Word document
    document = Document()
    document.add_heading(post.title, level=1)
    document.add_paragraph(markdown_content)
    
    # Save the document to a BytesIO buffer
    buffer = BytesIO()
    document.save(buffer)

    # Create a response with the Word document
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=post_document.docx'
    return response


def markdown(content):
    markdowner = Markdown()
    return markdowner.convert(content)